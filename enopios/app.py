import copy
import os
from inference.engine import Model
import copy
import streamlit as st
import mammoth
import docx
from io import BytesIO
import streamlit.components.v1 as components
from docx import Document

def generate_docx(array_content):
    # Create a DOCX file from the array content
    doc = docx.Document()
    for line in array_content:
        doc.add_paragraph(line)

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

def extract_text_from_docx(uploaded_file):
    result = mammoth.extract_raw_text(uploaded_file)
    text = result.value.splitlines()
    # Filter out empty lines
    #text = [line for line in text if line.strip()]

    return text 

def save_text_to_docx(text, file_path):
    doc = docx.Document()
    for line in text:
        doc.add_paragraph(line)
    
    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def translate_sentences(sentencess, src_lang, tgt_lang):
    sentences = copy.deepcopy(sentencess)
    print("======> translating")
    if src_lang == "hin_Deva":
        model_folder = "./ct2_model/indic-en-deploy/ct2_int8_model"
    else:
        model_folder = "./ct2_model/en-indic-deploy/ct2_int8_model"
   
    print("======> Translating ", src_lang, tgt_lang, model_folder)
    model = Model(model_folder, device="cuda", model_type="ctranslate2")
    # Perform translation using your inference engine model
    # Assuming 'batch_translate' is a function that performs batch translation
    #en_translations = model.batch_translate(sentences, src_lang, tgt_lang)
    #return en_translations

    max_tokens = 1024  # Maximum tokens per translation

    def split_sentences(sentences, max_tokens):
        split_sentences = []
        for sentence in sentences:
            if len(sentence.split()) > max_tokens:
                # Split the sentence into smaller chunks
                words = sentence.split()
                chunks = [words[i:i + max_tokens] for i in range(0, len(words), max_tokens)]
                for chunk in chunks:
                    split_sentences.append(' '.join(chunk))
            else:
                split_sentences.append(sentence)
        return split_sentences

    def join_translations(translations, original_sentences):
        joined_translations = []
        sentence_index = 0
        for sentence in original_sentences:
            if len(sentence.split()) > max_tokens:
                # Concatenate translations of split chunks
                joined_translation = ''
                words = sentence.split()
                while sentence_index < len(translations) and len(words) > max_tokens:
                    joined_translation += translations[sentence_index] + ' '
                    sentence_index += 1
                    words = words[max_tokens:]  # Move to the next chunk of words
                joined_translations.append(joined_translation.strip())
            else:
                joined_translations.append(translations[sentence_index])
                sentence_index += 1

        return joined_translations
    
    #Index the array to maintain original positions
    indexed_sentences = [(i, sentence) for i, sentence in enumerate(sentences)]
    
    # Filter out non-sentence elements (spaces, punctuation, etc.)
    only_sentences = [sentence for index, sentence in indexed_sentences if sentence.strip()]  # Adjust this based on your criteria

    # Split long sentences into chunks for translation
    split_chunks = split_sentences(only_sentences, max_tokens)

    # Translate the split chunks
    translated_chunks = model.batch_translate(split_chunks, src_lang, tgt_lang)
    
    # Join translated chunks back to their original sentences
    final_translations = join_translations(translated_chunks, only_sentences)
    
    final_translations.reverse()

    for i in range(len(sentences)): 
        s = sentences[i]
        if s.strip():
            sentences[i] = final_translations.pop() 

    return sentences 

def main():
    st.set_page_config(layout="wide")
    # Display external CSS for the PDF viewer
    
    st.markdown(
        '''
        <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet" integrity="sha384-T3c6CoIi6uLrA9TneNEoa7RxnatzjcDSCmG1MXxSR1GAsXEV/Dwwykc2MPK8M2HN" crossorigin="anonymous">
        ''',
        unsafe_allow_html=True
    )

    # Display PDF.js scripts
    st.markdown(
        '''
        <script type="module">
            import * as pdfjs from 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/4.0.269/pdf.min.mjs';
            pdfjs.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/4.0.269/pdf.worker.min.js';
        </script>
        ''',
        unsafe_allow_html=True
    )
    
    st.title("Document Translator")

    languages = {
        "Hindi to English": ("hin_Deva", "eng_Latn"),
        "English to Hindi": ("eng_Latn", "hin_Deva")
    }

    # Set default index to Hindi to English (index 0)
    default_language = "Hindi to English"
    default_index = 0

    option = st.selectbox("Select translation direction:", list(languages.keys()), index=default_index, format_func=lambda x: default_language if x == default_language else x)

    selected_lang = languages[option]
    src_lang = selected_lang[0]
    tgt_lang = selected_lang[1]
    
    uploaded_file = st.file_uploader("Upload DOCX file", type=["docx"])
    if uploaded_file is not None:
        # Create a temporary directory to store files
        temp_dir = os.getcwd() 

        
        pdf_file_path = "temp.pdf"
        docx_file_path = "temp.docx"
        # Check the file type of the uploaded file
        if uploaded_file.type == "application/pdf":
            # Convert the PDF file to DOCX format
            with open(pdf_file_path, "wb") as f:
                f.write(uploaded_file.getvalue())
            parse(pdf_file_path, docx_file_path)
            uploaded_file = open(docx_file_path, "rb")

        original_text = extract_text_from_docx(uploaded_file)

        print(st.session_state)
        # Check if translations already exist in the session state
        if 'translation_done' not in st.session_state or not st.session_state.translation_done:
            start_time = time.time()
            translated_text = translate_sentences(original_text, src_lang, tgt_lang)
            end_time = time.time()
            execution_time = end_time - start_time
            print(f"Execution time: {execution_time} seconds")
            # Save the translations to the session state
            st.session_state.translated_text = translated_text
            st.session_state.translation_done = True
        else:
            # If translations already exist in the session state, use them
            translated_text = st.session_state.translated_text
            
        html_content = ""
        for index, (original, translated) in enumerate(zip(original_text, translated_text)):
            if translated.strip():
                html_content += f'''
                    <div class="row" style="border-bottom: 1px solid #ccc; padding-bottom: 10px; margin-bottom: 10px;">
        <div class="col" style="align-self: flex-start;">
            <p style="margin-bottom: 0;">{original}</p>
        </div>
        <div class="col d-flex align-items-start">
            <textarea class="form-control h-100" readonly>{translated}</textarea>
        </div>
    </div>
                '''
            else:
                html_content += f'''
            <textarea class="form-control" rows="6">{translated}</textarea>
                '''

        # Combine all HTML content
        full_html_content = f'''
                <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet" integrity="sha384-T3c6CoIi6uLrA9TneNEoa7RxnatzjcDSCmG1MXxSR1GAsXEV/Dwwykc2MPK8M2HN" crossorigin="anonymous">
                <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/js/bootstrap.bundle.min.js" integrity="sha384-C6RzsynM9kWDrMNeT87bh95OGNyZPhcTNXj1NW7RuBCsyN/o0jlpcV8Qyq46cDfL" crossorigin="anonymous"></script>
               <center><div style="margin-top: 20px;">
    <a href="" class="btn btn-primary" id="download_docx">Download</a>
</div></center>
            <div class="mt-2 container">
                {html_content}
            </div>
                    <script>
                        document.getElementById('download_docx').addEventListener('click', function(e) {{
                            e.preventDefault();
                            var textToSave = '';
                            var textAreas = document.getElementsByTagName('textarea');
                            for (var i = 0; i < textAreas.length; i++) {{
                                textToSave += textAreas[i].value + '\\n';
                            }}

                            var blob = new Blob([textToSave], {{type: "text/plain;charset=utf-8"}});
                            var link = document.createElement('a');
                            link.href = window.URL.createObjectURL(blob);
                            link.download = "translated_text.docx";
                            link.click();
                        }});
                    </script>
        '''

        # Render HTML using st.components.v1.html
        st.components.v1.html(full_html_content, height=600, scrolling=True)

if __name__ == '__main__':
    main()
